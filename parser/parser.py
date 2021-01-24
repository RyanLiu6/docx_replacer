import os
import docx
import logging

from random import randint


class Parser:
    def __init__(self, file_path):
        """
        File path must be absolute
        """
        try:
            self.document = docx.Document(file_path)
        except FileNotFoundError:
            logging.error(f"Missing file at {file_path}")

        # set up some pathing stuff
        self.base_dir = os.path.abspath(os.getcwd())

    def paragraph_replace(self, old, new):
        for para in self.document.paragraphs:
            if old in para.text:
                para.text = para.text.replace(old, new)

        self._save_document()

    def table_replace(self, old, new):
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old in para.text:
                            print(para.text)

    def heading_replace(self, old, new):
        for content in self.document.paragraphs:
            if content.style.name=='Heading 1' or content.style.name=='Heading 2' or content.style.name=='Heading 3':
                print(content.text)

    def _save_document(self):
        random_name = f"{self.random_with_N_digits(6)}.docx"
        new_file_path = os.path.join(self.base_dir, f"output/{random_name}")

        with open(new_file_path, "w") as write_file:
            self.document.save(new_file_path)

        logging.info(f"New file written to {new_file_path}")

    def random_with_N_digits(self, n):
        range_start = 10**(n-1)
        range_end = (10**n)-1
        return randint(range_start, range_end)
