import os
import docx
import logging

from lib.utils import *

class Parser:
    def __init__(self, file_path):
        """
        File path must be absolute
        """
        try:
            self.document = docx.Document(file_path)
        except FileNotFoundError:
            logging.error(f"Missing file at {file_path}")

        # set up some pathing stuff
        self.base_dir = os.path.abspath(os.getcwd())

    def replace(self, old, new, input_type):
        if input_type == replace_types.every:
            self._replace_all(
                old=old,
                new=new)
        elif input_type == replace_types.tables:
            self._replace_tables(
                old=old,
                new=new)
        elif input_type == replace_types.headings:
            self._replace_headings(
                old=old,
                new=new)
        else:
            self._replace_paragraphs(
                old=old,
                new=new)

    def _replace_paragraphs(self, old, new, save=True):
        for para in self.document.paragraphs:
            if old in para.text:
                para.text = para.text.replace(old, new)

        if save:
            self._save_document()

    def _replace_tables(self, old, new, save=True):
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old in para.text:
                            para.text = para.text.replace(old, new)

        if save:
            self._save_document()

    def _replace_headings(self, old, new, save=True):
        for content in self.document.paragraphs:
            if content.style.name=='Heading 1' or content.style.name=='Heading 2' or content.style.name=='Heading 3':
                content.text = content.text.replace(old, new)

        if save:
            self._save_document()

    def _replace_all(self, old, new):
        self.paragraph_replace(
            old=old,
            new=new,
            save=False)

        self.table_replace(
            old=old,
            new=new,
            save=False)

        self.heading_replace(
            old=old,
            new=new,
            save=False)

        self._save_document()

    def _save_document(self):
        random_name = f"{random_with_N_digits(6)}.docx"
        new_file_path = os.path.join(self.base_dir, f"output/{random_name}")

        with open(new_file_path, "w") as write_file:
            self.document.save(new_file_path)

        logging.info(f"New file written to {new_file_path}")
